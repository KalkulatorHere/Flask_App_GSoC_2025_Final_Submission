"""
=============================================================================
 HYBRID OCR PIPELINE FOR 17TH-CENTURY SPANISH HISTORICAL DOCUMENTS
 Kaggle-Ready | End-to-End | Zero Fine-Tuning
=============================================================================

Pipeline Architecture (5 Stages):
  1. Preprocessing      → Sauvola binarization + deskewing       [RESEARCH]
  2. Layout Detection   → Detectron2 Mask R-CNN R101-FPN          [EXISTING CODEBASE]
  3. Line Segmentation  → Kraken                                  [RESEARCH]
  4. OCR Recognition    → qantev/trocr-large-spanish (TrOCR)      [EXISTING CODEBASE]
  5. LLM/VLM Cleanup    → GPT-4o + Gemini two-pass                [RESEARCH]

Architecture: Transformer-based (TrOCR = ViT encoder + language model decoder)
Fine-tuning: ZERO — all models are pre-trained

Evaluation Metrics:
  - CER (Character Error Rate) — primary
  - WER (Word Error Rate)      — secondary
  - F1 (Layout detection)      — for Stage 2
  - HCPR (Historical Char Preservation Rate) — for archaic chars
"""

# ============================================================================
# SECTION 0: INSTALLATION (Kaggle cell — run first)
# ============================================================================

import subprocess
import sys

def install_packages():
    """Install all required packages for Kaggle environment."""
    packages = [
        "torch", "torchvision",
        "transformers",
        "opencv-python-headless",
        "Pillow",
        "numpy",
        "editdistance",
        "jiwer",
        "google-generativeai",
        "openai",
        "python-docx",
        "PyMuPDF",
    ]
    
    # Detectron2 (special install for Kaggle)
    print("=" * 60)
    print("Installing packages...")
    print("=" * 60)
    
    for pkg in packages:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", pkg])
    
    # Detectron2 install
    try:
        import detectron2
        print("✅ detectron2 already installed")
    except ImportError:
        print("Installing detectron2...")
        subprocess.check_call([
            sys.executable, "-m", "pip", "install", "-q",
            "detectron2", "-f",
            "https://dl.fbaipublicfiles.com/detectron2/wheels/cu118/torch2.1/index.html"
        ])
    
    # Kraken install
    try:
        import kraken
        print("✅ kraken already installed")
    except ImportError:
        print("Installing kraken...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", "kraken"])
    
    print("✅ All packages installed!")

# Uncomment the line below to install packages on Kaggle:
# install_packages()


# ============================================================================
# SECTION 1: IMPORTS
# ============================================================================

import os
import logging
import time
import json
import glob
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass, field

import numpy as np
import cv2
from PIL import Image
import torch

# TrOCR
from transformers import TrOCRProcessor, VisionEncoderDecoderModel

# Detectron2
from detectron2.config import get_cfg
from detectron2.engine import DefaultPredictor
from detectron2.data import MetadataCatalog
from detectron2 import model_zoo

# Kraken line segmentation
try:
    from kraken import blla
    from kraken.lib import vgsl
    from kraken.containers import Segmentation
    KRAKEN_AVAILABLE = True
    print("✅ Kraken available")
except ImportError:
    KRAKEN_AVAILABLE = False
    print("⚠️ Kraken not available — will use fallback line segmentation")

# Evaluation
import editdistance
try:
    from jiwer import wer as compute_wer_jiwer
    JIWER_AVAILABLE = True
except ImportError:
    JIWER_AVAILABLE = False

# LLM APIs
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    print("⚠️ Gemini not available")

try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print("⚠️ OpenAI not available")

# PyMuPDF for PDF rendering
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
    print("✅ PyMuPDF available")
except ImportError:
    FITZ_AVAILABLE = False
    print("⚠️ PyMuPDF not available — install with: pip install PyMuPDF")

# Docx for ground truth
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not available — install with: pip install python-docx")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("OCR_Pipeline")


# ============================================================================
# SECTION 2: CONFIGURATION
# ============================================================================

@dataclass
class PipelineConfig:
    """Configuration for the entire OCR pipeline."""
    
    # --- Paths ---
    # UPDATE THESE PATHS for your Kaggle environment!
    detectron2_weights: str = "/kaggle/input/your-dataset/model_final (8) (1).pth"
    input_pdf: str = "/kaggle/input/your-dataset/document.pdf"
    output_dir: str = "/kaggle/working/ocr_output/"
    ground_truth_docx: str = "/kaggle/input/your-dataset/ground_truth.docx"
    pdf_dpi: int = 200  # DPI for PDF rendering
    
    # --- Detectron2 (Stage 2) ---
    detectron2_config: str = "COCO-InstanceSegmentation/mask_rcnn_R_101_FPN_3x.yaml"
    detectron2_num_classes: int = 2  # textline, baseline
    detectron2_score_threshold: float = 0.5
    
    # --- Embellishment filtering ---
    area_filter_threshold_percent: float = 12.5  # boxes below this % of avg area → discard
    
    # --- TrOCR (Stage 4) ---
    trocr_primary_model: str = "qantev/trocr-large-spanish"
    trocr_fallback_model: str = "microsoft/trocr-base-printed"
    trocr_max_new_tokens: int = 128
    
    # --- LLM APIs (Stage 5) ---
    openai_api_key: str = ""   # Set via env: OPENAI_API_KEY
    gemini_api_key: str = ""   # Set via env: GEMINI_API_KEY
    gemini_model: str = "gemini-2.5-flash-preview-05-20"
    openai_model: str = "gpt-4o"
    
    # --- Preprocessing (Stage 1) ---
    target_dpi: int = 300
    sauvola_window_size: int = 25
    sauvola_k: float = 0.15
    
    # --- Device ---
    device: str = "auto"  # "auto", "cuda", "cpu"
    
    def __post_init__(self):
        os.makedirs(self.output_dir, exist_ok=True)
        
        if not self.openai_api_key:
            self.openai_api_key = os.getenv("OPENAI_API_KEY", "")
        if not self.gemini_api_key:
            self.gemini_api_key = os.getenv("GEMINI_API_KEY", "")


# ============================================================================
# SECTION 3: STAGE 1 — PREPROCESSING (from Research)
# ============================================================================

class Preprocessor:
    """Stage 1: Image preprocessing with Sauvola binarization and deskewing."""
    
    def __init__(self, config: PipelineConfig):
        self.config = config
    
    def sauvola_binarize(self, image: np.ndarray) -> np.ndarray:
        """Apply Sauvola adaptive thresholding for binarization."""
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image.copy()
        
        # Sauvola thresholding (adaptive)
        binary = cv2.adaptiveThreshold(
            gray, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            self.config.sauvola_window_size,
            self.config.sauvola_k * 100  # OpenCV uses C parameter differently
        )
        return binary
    
    def deskew(self, image: np.ndarray) -> np.ndarray:
        """Correct page skew using minimum area rectangle."""
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image.copy()
        
        # Find contours
        thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
        coords = np.column_stack(np.where(thresh > 0))
        
        if len(coords) < 100:
            return image
        
        # Get minimum area rectangle angle
        angle = cv2.minAreaRect(coords)[-1]
        
        if angle < -45:
            angle = -(90 + angle)
        else:
            angle = -angle
        
        # Only deskew if angle is small (avoid flipping)
        if abs(angle) > 5:
            return image
        
        if abs(angle) < 0.5:
            return image
        
        # Rotate
        h, w = image.shape[:2]
        center = (w // 2, h // 2)
        M = cv2.getRotationMatrix2D(center, angle, 1.0)
        rotated = cv2.warpAffine(image, M, (w, h), flags=cv2.INTER_CUBIC,
                                  borderMode=cv2.BORDER_REPLICATE)
        return rotated
    
    def denoise(self, image: np.ndarray) -> np.ndarray:
        """Remove noise using non-local means denoising."""
        if len(image.shape) == 3:
            return cv2.fastNlMeansDenoisingColored(image, None, 10, 10, 7, 21)
        else:
            return cv2.fastNlMeansDenoising(image, None, 10, 7, 21)
    
    def preprocess(self, image: np.ndarray) -> Tuple[np.ndarray, np.ndarray]:
        """
        Full preprocessing pipeline.
        Returns: (preprocessed_color, binarized_gray)
        """
        logger.info("Stage 1: Preprocessing...")
        
        # Denoise
        denoised = self.denoise(image)
        
        # Deskew
        deskewed = self.deskew(denoised)
        
        # Binarize (for display/debugging)
        binarized = self.sauvola_binarize(deskewed)
        
        logger.info("  ✅ Preprocessing complete (denoise → deskew → binarize)")
        return deskewed, binarized


# ============================================================================
# SECTION 4: STAGE 2 — LAYOUT & TEXT DETECTION (from Existing Codebase)
# ============================================================================

class TextlineDetector:
    """
    Stage 2: Textline detection using Detectron2 Mask R-CNN.
    From existing codebase — NO fine-tuning needed.
    Uses pre-trained weights with area-based embellishment filtering.
    """
    
    def __init__(self, config: PipelineConfig):
        self.config = config
        self.predictor = None
        self.device = self._get_device()
        self._setup_and_load()
    
    def _get_device(self) -> torch.device:
        if self.config.device == "auto":
            return torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        return torch.device(self.config.device)
    
    def _setup_and_load(self):
        """Setup Detectron2 config and load pre-trained weights."""
        logger.info("Stage 2: Loading Detectron2 textline detection model...")
        
        cfg = get_cfg()
        cfg.merge_from_file(model_zoo.get_config_file(self.config.detectron2_config))
        cfg.MODEL.ROI_HEADS.NUM_CLASSES = self.config.detectron2_num_classes
        cfg.MODEL.ROI_HEADS.SCORE_THRESH_TEST = self.config.detectron2_score_threshold
        cfg.MODEL.WEIGHTS = self.config.detectron2_weights
        cfg.DATASETS.TEST = ("page_test",)
        cfg.DATALOADER.NUM_WORKERS = 2
        
        # Set device
        if not torch.cuda.is_available():
            cfg.MODEL.DEVICE = "cpu"
        
        MetadataCatalog.get("page_test").thing_classes = ["textline", "baseline"]
        
        self.predictor = DefaultPredictor(cfg)
        logger.info(f"  ✅ Detectron2 model loaded on {self.device}")
    
    def detect_textlines(self, image: np.ndarray) -> Tuple[np.ndarray, np.ndarray]:
        """Detect textline bounding boxes in image."""
        outputs = self.predictor(image)
        instances = outputs["instances"].to("cpu")
        
        # Filter for textline class (class 0 = textline)
        textline_mask = instances.pred_classes == 0
        boxes = instances.pred_boxes[textline_mask].tensor.numpy()
        scores = instances.scores[textline_mask].numpy()
        
        logger.info(f"  Detected {len(boxes)} raw textlines")
        return boxes, scores
    
    def filter_embellishments(self, boxes: np.ndarray, scores: np.ndarray
                               ) -> Tuple[np.ndarray, np.ndarray, np.ndarray, np.ndarray]:
        """
        Filter out embellishments using area-based heuristic.
        Boxes with area < threshold% of average → embellishments → discard.
        """
        if len(boxes) == 0:
            return np.array([]), np.array([]), np.array([]), np.array([])
        
        areas = np.array([(b[2]-b[0]) * (b[3]-b[1]) for b in boxes])
        avg_area = np.mean(areas)
        threshold = avg_area * (self.config.area_filter_threshold_percent / 100.0)
        
        mask = areas >= threshold
        main_boxes = boxes[mask]
        main_scores = scores[mask]
        margin_boxes = boxes[~mask]
        margin_scores = scores[~mask]
        
        logger.info(f"  Kept {len(main_boxes)} main textlines, discarded {len(margin_boxes)} embellishments")
        return main_boxes, main_scores, margin_boxes, margin_scores
    
    def sort_reading_order(self, boxes: np.ndarray, scores: np.ndarray
                           ) -> Tuple[np.ndarray, np.ndarray, List[Dict]]:
        """Sort textlines top-to-bottom by y-coordinate."""
        if len(boxes) == 0:
            return boxes, scores, []
        
        centers_y = np.array([(b[1] + b[3]) / 2 for b in boxes])
        sort_indices = np.argsort(centers_y)
        
        sorted_boxes = boxes[sort_indices]
        sorted_scores = scores[sort_indices]
        reading_order = []
        
        for pos, idx in enumerate(sort_indices):
            reading_order.append({
                'original_index': int(idx),
                'column': 0,
                'position_in_column': int(pos),
                'reading_order_index': int(pos)
            })
        
        return sorted_boxes, sorted_scores, reading_order
    
    def calculate_dynamic_padding(self, boxes: np.ndarray, image_shape: Tuple
                                   ) -> Dict[str, int]:
        """Calculate adaptive padding based on inter-box distances."""
        if len(boxes) < 2:
            return {"top": 10, "bottom": 10, "left": 8, "right": 8}
        
        # Sort by y
        centers_y = np.array([(b[1]+b[3])/2 for b in boxes])
        sorted_idx = np.argsort(centers_y)
        sorted_boxes = boxes[sorted_idx]
        
        # Vertical gaps
        v_gaps = []
        for i in range(len(sorted_boxes) - 1):
            cx1 = (sorted_boxes[i][0] + sorted_boxes[i][2]) / 2
            cx2 = (sorted_boxes[i+1][0] + sorted_boxes[i+1][2]) / 2
            if abs(cx1 - cx2) < image_shape[1] * 0.3:
                gap = sorted_boxes[i+1][1] - sorted_boxes[i][3]
                if gap > 0:
                    v_gaps.append(gap)
        
        avg_vgap = np.median(v_gaps) if v_gaps else 20
        v_pad = max(5, min(25, avg_vgap / 2))
        h_pad = max(3, min(20, 15 / 3))
        
        # Height factor
        heights = [b[3]-b[1] for b in boxes]
        avg_h = np.mean(heights)
        h_factor = max(0.1, min(0.3, avg_h / 100))
        v_pad = max(v_pad, avg_h * h_factor)
        
        return {
            "top": int(v_pad * 0.95),
            "bottom": int(v_pad * 1.2),
            "left": int(h_pad),
            "right": int(h_pad)
        }
    
    def crop_textlines(self, image: np.ndarray, boxes: np.ndarray
                       ) -> Tuple[List[np.ndarray], List[List[int]], Dict]:
        """Crop textline regions with dynamic padding."""
        if len(boxes) == 0:
            return [], [], {}
        
        padding = self.calculate_dynamic_padding(boxes, image.shape)
        crops = []
        padded_boxes = []
        
        for box in boxes:
            x1, y1, x2, y2 = box.astype(int)
            x1p = max(0, x1 - padding["left"])
            y1p = max(0, y1 - padding["top"])
            x2p = min(image.shape[1], x2 + padding["right"])
            y2p = min(image.shape[0], y2 + padding["bottom"])
            
            crop = image[y1p:y2p, x1p:x2p]
            if crop.size > 0:
                crops.append(crop)
                padded_boxes.append([x1p, y1p, x2p, y2p])
        
        return crops, padded_boxes, padding
    
    def detect_and_crop(self, image: np.ndarray) -> Dict:
        """Full Stage 2: detect → filter → sort → crop."""
        logger.info("Stage 2: Layout & Text Detection...")
        
        boxes, scores = self.detect_textlines(image)
        if len(boxes) == 0:
            return {'success': False, 'error': 'No textlines detected'}
        
        main_boxes, main_scores, margin_boxes, margin_scores = \
            self.filter_embellishments(boxes, scores)
        
        if len(main_boxes) == 0:
            return {'success': False, 'error': 'No textlines after filtering'}
        
        ordered_boxes, ordered_scores, reading_order = \
            self.sort_reading_order(main_boxes, main_scores)
        
        crops, padded_boxes, padding = self.crop_textlines(image, ordered_boxes)
        
        logger.info(f"  ✅ Layout detection complete: {len(crops)} text regions")
        
        return {
            'success': True,
            'crops': crops,
            'boxes': ordered_boxes,
            'scores': ordered_scores,
            'padded_boxes': padded_boxes,
            'reading_order': reading_order,
            'padding': padding,
            'margin_boxes': margin_boxes,
            'total_detected': len(boxes),
            'total_filtered': len(margin_boxes)
        }


# ============================================================================
# SECTION 5: STAGE 3 — LINE SEGMENTATION (from Research — Kraken)
# ============================================================================

class LineSegmenter:
    """
    Stage 3: Line segmentation using Kraken.
    Falls back to direct crop usage if Kraken is not available.
    """
    
    def __init__(self, config: PipelineConfig):
        self.config = config
    
    def segment_block(self, block_image: np.ndarray) -> List[np.ndarray]:
        """
        Segment a text block into individual lines using Kraken.
        If Kraken unavailable, returns the block as a single 'line'.
        """
        if not KRAKEN_AVAILABLE:
            return [block_image]
        
        try:
            # Convert to PIL
            if len(block_image.shape) == 3:
                pil_img = Image.fromarray(cv2.cvtColor(block_image, cv2.COLOR_BGR2RGB))
            else:
                pil_img = Image.fromarray(block_image)
            
            # Use Kraken baseline segmentation
            seg_result = blla.segment(pil_img)
            
            lines = []
            if hasattr(seg_result, 'lines') and len(seg_result.lines) > 0:
                for line in seg_result.lines:
                    # Extract line crop using bounding box from baseline
                    if hasattr(line, 'bbox'):
                        x1, y1, x2, y2 = line.bbox
                        # Add small padding
                        pad = 3
                        x1 = max(0, x1 - pad)
                        y1 = max(0, y1 - pad)
                        x2 = min(block_image.shape[1], x2 + pad)
                        y2 = min(block_image.shape[0], y2 + pad)
                        crop = block_image[y1:y2, x1:x2]
                        if crop.size > 0:
                            lines.append(crop)
                
                if len(lines) > 0:
                    return lines
            
            # If Kraken returned nothing, return the whole block
            return [block_image]
            
        except Exception as e:
            logger.warning(f"  Kraken segmentation failed: {e}, using block as single line")
            return [block_image]
    
    def segment_all_blocks(self, crops: List[np.ndarray]) -> List[np.ndarray]:
        """
        Segment all text block crops into individual lines.
        If Kraken works well, each block may produce multiple lines.
        If Kraken fails, each block is treated as one line (Detectron2 already 
        detects individual textlines, so this is fine).
        """
        logger.info("Stage 3: Line Segmentation (Kraken)...")
        
        all_lines = []
        for i, crop in enumerate(crops):
            lines = self.segment_block(crop)
            all_lines.extend(lines)
        
        if KRAKEN_AVAILABLE:
            logger.info(f"  ✅ Kraken segmented {len(crops)} blocks → {len(all_lines)} lines")
        else:
            logger.info(f"  ⚠️ Kraken unavailable — using {len(all_lines)} Detectron2 crops as lines")
        
        return all_lines


# ============================================================================
# SECTION 6: STAGE 4 — OCR RECOGNITION (from Existing Codebase — TrOCR)
# ============================================================================

class OCRRecognizer:
    """
    Stage 4: OCR using pre-trained TrOCR Spanish model.
    From existing codebase — NO fine-tuning needed.
    Model: qantev/trocr-large-spanish (fallback: microsoft/trocr-base-printed)
    Architecture: Transformer (ViT encoder → Language Model decoder)
    """
    
    def __init__(self, config: PipelineConfig):
        self.config = config
        self.processor = None
        self.model = None
        self.device = self._get_device()
        self.model_name = ""
        self._load_model()
    
    def _get_device(self) -> torch.device:
        if self.config.device == "auto":
            return torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        return torch.device(self.config.device)
    
    def _load_model(self):
        """Load TrOCR model with Spanish primary + English fallback."""
        logger.info("Stage 4: Loading TrOCR model...")
        
        try:
            self.processor = TrOCRProcessor.from_pretrained(self.config.trocr_primary_model)
            self.model = VisionEncoderDecoderModel.from_pretrained(self.config.trocr_primary_model)
            self.model_name = self.config.trocr_primary_model
            logger.info(f"  ✅ Primary model loaded: {self.config.trocr_primary_model}")
        except Exception as e:
            logger.warning(f"  ⚠️ Spanish model failed ({e}), loading fallback...")
            self.processor = TrOCRProcessor.from_pretrained(self.config.trocr_fallback_model)
            self.model = VisionEncoderDecoderModel.from_pretrained(self.config.trocr_fallback_model)
            self.model_name = self.config.trocr_fallback_model
            logger.info(f"  ✅ Fallback model loaded: {self.config.trocr_fallback_model}")
        
        self.model.to(self.device)
        logger.info(f"  TrOCR on device: {self.device}")
    
    def recognize_line(self, line_image: np.ndarray) -> Dict:
        """Recognize text from a single line image."""
        try:
            # Convert to PIL RGB
            if len(line_image.shape) == 2:
                pil_img = Image.fromarray(line_image).convert('RGB')
            else:
                pil_img = Image.fromarray(cv2.cvtColor(line_image, cv2.COLOR_BGR2RGB))
            
            # TrOCR inference
            pixel_values = self.processor(images=pil_img, return_tensors="pt").pixel_values
            pixel_values = pixel_values.to(self.device)
            
            with torch.no_grad():
                generated_ids = self.model.generate(
                    pixel_values,
                    max_new_tokens=self.config.trocr_max_new_tokens
                )
                text = self.processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
            
            return {'text': text.strip(), 'confidence': 1.0, 'success': True}
            
        except Exception as e:
            logger.error(f"  OCR error: {e}")
            return {'text': '', 'confidence': 0.0, 'success': False, 'error': str(e)}
    
    def recognize_all(self, line_images: List[np.ndarray]) -> List[Dict]:
        """Recognize text from all line images."""
        logger.info(f"Stage 4: OCR Recognition ({len(line_images)} lines)...")
        
        results = []
        for i, img in enumerate(line_images):
            result = self.recognize_line(img)
            result['line_index'] = i
            results.append(result)
            
            if (i + 1) % 10 == 0:
                logger.info(f"  Processed {i+1}/{len(line_images)} lines...")
        
        successful = sum(1 for r in results if r['success'])
        logger.info(f"  ✅ OCR complete: {successful}/{len(results)} lines recognized")
        
        return results


# ============================================================================
# SECTION 7: STAGE 5 — LLM/VLM POST-PROCESSING (from Research)
# ============================================================================

class LLMPostProcessor:
    """
    Stage 5: Two-pass LLM/VLM post-processing.
    Pass 1: GPT-4o text correction (always runs)
    Pass 2: Gemini visual verification (for low-confidence pages)
    Zero fine-tuning — uses zero-shot prompting.
    """
    
    CORRECTION_PROMPT = """You are an expert in 17th-century Spanish legal documents. 
Given OCR output from a historical printed source, correct errors while following these rules:

1. Fix character-level OCR errors (e.g., 'rn' misread as 'm', 'cl' as 'd')
2. Fix word segmentation errors (merged or split words)
3. PRESERVE historical spelling (e.g., 'honefta' NOT 'honesta', 'defhonra' NOT 'deshonra')
4. PRESERVE abbreviations and archaic characters
5. DO NOT modernize the language or grammar
6. Fix obvious punctuation OCR errors
7. Maintain original line structure where possible

Return ONLY the corrected text, nothing else.

OCR Text to correct:
{text}"""

    VISUAL_VERIFY_PROMPT = """Here is a page image from a 17th-century Spanish legal document 
and its current OCR transcription. Please:

1. Verify the transcription against the visible text in the image
2. Correct any remaining errors you can see
3. Identify any text the OCR may have missed
4. IGNORE decorative elements — only transcribe main body text
5. PRESERVE historical spelling exactly as printed

Current transcription:
{text}

Return ONLY the verified/corrected text."""
    
    def __init__(self, config: PipelineConfig):
        self.config = config
        self._setup_apis()
    
    def _setup_apis(self):
        """Configure LLM API clients."""
        # OpenAI
        self.openai_client = None
        if OPENAI_AVAILABLE and self.config.openai_api_key:
            self.openai_client = openai.OpenAI(api_key=self.config.openai_api_key)
            logger.info("  ✅ OpenAI API configured")
        
        # Gemini
        self.gemini_available = False
        if GEMINI_AVAILABLE and self.config.gemini_api_key:
            genai.configure(api_key=self.config.gemini_api_key)
            self.gemini_available = True
            logger.info("  ✅ Gemini API configured")
    
    def pass1_gpt4o_correction(self, raw_text: str) -> Tuple[str, str]:
        """
        Pass 1: GPT-4o text-only correction.
        Falls back to Gemini text correction if OpenAI unavailable.
        """
        prompt = self.CORRECTION_PROMPT.format(text=raw_text)
        
        # Try OpenAI GPT-4o
        if self.openai_client:
            try:
                response = self.openai_client.chat.completions.create(
                    model=self.config.openai_model,
                    messages=[
                        {"role": "system", "content": "You are an expert transcriber of 17th-century Spanish legal documents."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.1,
                    max_tokens=4096
                )
                corrected = response.choices[0].message.content.strip()
                return corrected, "gpt-4o"
            except Exception as e:
                logger.warning(f"  GPT-4o correction failed: {e}")
        
        # Fallback to Gemini for text correction
        if self.gemini_available:
            try:
                model = genai.GenerativeModel(self.config.gemini_model)
                response = model.generate_content(prompt)
                if response.candidates and response.text:
                    return response.text.strip(), "gemini_text"
            except Exception as e:
                logger.warning(f"  Gemini text correction failed: {e}")
        
        logger.warning("  ⚠️ No LLM available for Pass 1, returning raw text")
        return raw_text, "none"
    
    def pass2_gemini_visual_verify(self, image: np.ndarray, corrected_text: str) -> Tuple[str, str]:
        """
        Pass 2: Gemini visual verification.
        Sends page image + text to Gemini for cross-checking.
        """
        if not self.gemini_available:
            return corrected_text, "visual_skip"
        
        try:
            # Convert image to PIL
            if len(image.shape) == 3:
                pil_img = Image.fromarray(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
            else:
                pil_img = Image.fromarray(image)
            
            prompt = self.VISUAL_VERIFY_PROMPT.format(text=corrected_text)
            
            model = genai.GenerativeModel(self.config.gemini_model)
            response = model.generate_content([prompt, pil_img])
            
            if response.candidates and response.text:
                return response.text.strip(), "gemini_visual"
                
        except Exception as e:
            logger.warning(f"  Gemini visual verification failed: {e}")
        
        return corrected_text, "visual_failed"
    
    def process(self, raw_text: str, image: np.ndarray = None,
                run_visual_pass: bool = True) -> Dict:
        """
        Full Stage 5: Two-pass LLM/VLM post-processing.
        """
        logger.info("Stage 5: LLM/VLM Post-Processing...")
        
        # Pass 1: Text correction (always)
        logger.info("  Pass 1: LLM text correction...")
        corrected_text, pass1_backend = self.pass1_gpt4o_correction(raw_text)
        logger.info(f"  ✅ Pass 1 complete (backend: {pass1_backend})")
        
        # Pass 2: Visual verification (optional)
        final_text = corrected_text
        pass2_backend = "skipped"
        
        if run_visual_pass and image is not None:
            logger.info("  Pass 2: VLM visual verification...")
            final_text, pass2_backend = self.pass2_gemini_visual_verify(image, corrected_text)
            logger.info(f"  ✅ Pass 2 complete (backend: {pass2_backend})")
        
        return {
            'raw_text': raw_text,
            'pass1_corrected': corrected_text,
            'pass1_backend': pass1_backend,
            'final_text': final_text,
            'pass2_backend': pass2_backend,
        }


# ============================================================================
# SECTION 8: EVALUATION METRICS
# ============================================================================

class OCREvaluator:
    """
    Evaluation metrics for OCR pipeline.
    Primary: CER (Character Error Rate), WER (Word Error Rate)
    Secondary: F1 (Layout), HCPR (Historical Character Preservation Rate)
    """
    
    @staticmethod
    def compute_cer(prediction: str, reference: str) -> float:
        """Compute Character Error Rate using Levenshtein distance."""
        if len(reference) == 0:
            return 0.0 if len(prediction) == 0 else 1.0
        distance = editdistance.eval(prediction, reference)
        return distance / len(reference)
    
    @staticmethod
    def compute_wer(prediction: str, reference: str) -> float:
        """Compute Word Error Rate."""
        if JIWER_AVAILABLE:
            try:
                return compute_wer_jiwer(reference, prediction)
            except:
                pass
        
        # Manual WER
        ref_words = reference.split()
        pred_words = prediction.split()
        if len(ref_words) == 0:
            return 0.0 if len(pred_words) == 0 else 1.0
        distance = editdistance.eval(ref_words, pred_words)
        return distance / len(ref_words)
    
    @staticmethod
    def compute_hcpr(prediction: str, reference: str) -> float:
        """
        Compute Historical Character Preservation Rate.
        Measures if archaic characters are preserved (not modernized).
        """
        archaic_chars = set('ſꝯ&ꝑꝙ')  # long-s, special abbreviations
        
        ref_archaic = [c for c in reference if c in archaic_chars]
        if len(ref_archaic) == 0:
            return 1.0  # No archaic chars to check
        
        preserved = 0
        for i, c in enumerate(ref_archaic):
            if c in prediction:
                preserved += 1
        
        return preserved / len(ref_archaic)
    
    @staticmethod
    def compute_bow_accuracy(prediction: str, reference: str) -> float:
        """Compute Bag-of-Words accuracy (order-independent)."""
        ref_words = set(reference.split())
        pred_words = set(prediction.split())
        
        if len(ref_words) == 0:
            return 1.0 if len(pred_words) == 0 else 0.0
        
        overlap = ref_words & pred_words
        return len(overlap) / len(ref_words)
    
    def evaluate(self, predictions: List[str], references: List[str]) -> Dict:
        """Run all evaluation metrics."""
        assert len(predictions) == len(references), "Predictions and references must match"
        
        n = len(predictions)
        metrics = {
            'per_page': [],
            'aggregate': {}
        }
        
        total_cer, total_wer, total_hcpr, total_bow = 0, 0, 0, 0
        
        for i, (pred, ref) in enumerate(zip(predictions, references)):
            cer = self.compute_cer(pred, ref)
            w = self.compute_wer(pred, ref)
            hcpr = self.compute_hcpr(pred, ref)
            bow = self.compute_bow_accuracy(pred, ref)
            
            metrics['per_page'].append({
                'page': i,
                'cer': round(cer, 4),
                'wer': round(w, 4),
                'hcpr': round(hcpr, 4),
                'bow_accuracy': round(bow, 4)
            })
            
            total_cer += cer
            total_wer += w
            total_hcpr += hcpr
            total_bow += bow
        
        metrics['aggregate'] = {
            'avg_cer': round(total_cer / n, 4),
            'avg_wer': round(total_wer / n, 4),
            'avg_hcpr': round(total_hcpr / n, 4),
            'avg_bow_accuracy': round(total_bow / n, 4),
            'num_pages': n
        }
        
        return metrics


# ============================================================================
# SECTION 9: MAIN PIPELINE (End-to-End)
# ============================================================================

class HybridOCRPipeline:
    """
    End-to-end hybrid OCR pipeline for 17th-century Spanish documents.
    
    Architecture: Transformer-based (TrOCR)
    Fine-tuning: ZERO
    
    Stages:
      1. Preprocessing      → Sauvola + deskew        (Research)
      2. Layout Detection    → Detectron2 R-CNN        (Existing Codebase)
      3. Line Segmentation   → Kraken                  (Research)
      4. OCR Recognition     → TrOCR Spanish           (Existing Codebase)
      5. LLM/VLM Cleanup     → GPT-4o + Gemini         (Research)
    """
    
    def __init__(self, config: PipelineConfig):
        self.config = config
        
        print("=" * 60)
        print("  HYBRID OCR PIPELINE — Initializing")
        print("=" * 60)
        
        # Initialize all stages
        self.preprocessor = Preprocessor(config)
        self.textline_detector = TextlineDetector(config)
        self.line_segmenter = LineSegmenter(config)
        self.ocr_recognizer = OCRRecognizer(config)
        self.llm_processor = LLMPostProcessor(config)
        self.evaluator = OCREvaluator()
        
        print("=" * 60)
        print("  ✅ Pipeline ready!")
        print("=" * 60)
    
    def process_single_image(self, image_path: str,
                              run_visual_pass: bool = True,
                              save_intermediates: bool = True) -> Dict:
        """
        Process a single page image through the full pipeline.
        
        Args:
            image_path: Path to the page image
            run_visual_pass: Whether to run Gemini visual verification (Pass 2)
            save_intermediates: Whether to save intermediate outputs
        
        Returns:
            Dict with all pipeline outputs
        """
        print(f"\n{'='*60}")
        print(f"  Processing: {os.path.basename(image_path)}")
        print(f"{'='*60}")
        
        start_time = time.time()
        
        # Load image
        image = cv2.imread(image_path)
        if image is None:
            return {'success': False, 'error': f'Could not load image: {image_path}'}
        
        # -------- STAGE 1: Preprocessing (Research) --------
        preprocessed, binarized = self.preprocessor.preprocess(image)
        
        # -------- STAGE 2: Layout Detection (Existing Codebase) --------
        detection_result = self.textline_detector.detect_and_crop(preprocessed)
        if not detection_result['success']:
            return detection_result
        
        crops = detection_result['crops']
        
        # -------- STAGE 3: Line Segmentation (Research — Kraken) --------
        line_images = self.line_segmenter.segment_all_blocks(crops)
        
        # -------- STAGE 4: OCR Recognition (Existing Codebase — TrOCR) --------
        ocr_results = self.ocr_recognizer.recognize_all(line_images)
        
        # Assemble raw text
        raw_lines = [r['text'] for r in ocr_results if r['text'].strip()]
        raw_text = "\n".join(raw_lines)
        
        # -------- STAGE 5: LLM/VLM Post-Processing (Research) --------
        llm_result = self.llm_processor.process(
            raw_text, image=image, run_visual_pass=run_visual_pass
        )
        
        elapsed = time.time() - start_time
        
        # Save outputs
        result = {
            'success': True,
            'image_path': image_path,
            'raw_text': raw_text,
            'corrected_text': llm_result['pass1_corrected'],
            'final_text': llm_result['final_text'],
            'pass1_backend': llm_result['pass1_backend'],
            'pass2_backend': llm_result['pass2_backend'],
            'num_lines_detected': len(crops),
            'num_lines_segmented': len(line_images),
            'num_lines_recognized': sum(1 for r in ocr_results if r['success']),
            'total_detected_boxes': detection_result['total_detected'],
            'filtered_embellishments': detection_result['total_filtered'],
            'trocr_model': self.ocr_recognizer.model_name,
            'processing_time_seconds': round(elapsed, 2),
            'ocr_results_per_line': ocr_results,
        }
        
        if save_intermediates:
            self._save_outputs(result, image_path, line_images, binarized)
        
        print(f"\n  ✅ Done in {elapsed:.1f}s — {len(raw_lines)} lines recognized")
        print(f"  Raw text length: {len(raw_text)} chars")
        print(f"  Final text length: {len(llm_result['final_text'])} chars")
        
        return result
    
    def process_pdf(self, pdf_path: str = None,
                     run_visual_pass: bool = True,
                     dpi: int = None) -> Dict:
        """
        Process a PDF file: render each page to image, then run pipeline.
        
        Args:
            pdf_path: Path to the PDF file
            run_visual_pass: Whether to run Gemini visual verification
            dpi: DPI for PDF rendering (default: config.pdf_dpi)
        """
        if not FITZ_AVAILABLE:
            return {'success': False, 'error': 'PyMuPDF not installed. Run: pip install PyMuPDF'}
        
        if pdf_path is None:
            pdf_path = self.config.input_pdf
        
        if not os.path.exists(pdf_path):
            return {'success': False, 'error': f'PDF not found: {pdf_path}'}
        
        if dpi is None:
            dpi = self.config.pdf_dpi
        
        print(f"\n{'='*60}")
        print(f"  Loading PDF: {os.path.basename(pdf_path)}")
        print(f"{'='*60}")
        
        # Render PDF pages to images
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        print(f"  Total pages in PDF: {total_pages}")
        
        # Save rendered pages as images (with spread/double-page detection)
        pages_dir = os.path.join(self.config.output_dir, "rendered_pages")
        os.makedirs(pages_dir, exist_ok=True)
        
        image_files = []
        for i in range(total_pages):
            page = doc.load_page(i)
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            nparr = np.frombuffer(img_bytes, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            
            h, w = img.shape[:2]
            
            # Spread detection: if width > 1.2× height, this PDF page
            # contains two actual pages side by side → split into left + right
            if w > h * 1.2:
                mid = w // 2
                left_img = img[:, :mid]
                right_img = img[:, mid:]
                
                left_filename = f"page_{i+1:03d}a_left.png"
                right_filename = f"page_{i+1:03d}b_right.png"
                
                left_path = os.path.join(pages_dir, left_filename)
                right_path = os.path.join(pages_dir, right_filename)
                
                cv2.imwrite(left_path, left_img)
                cv2.imwrite(right_path, right_img)
                
                image_files.append(left_path)
                image_files.append(right_path)
                print(f"  Page {i+1}/{total_pages}: spread ({w}×{h}) → split into {left_filename} + {right_filename}")
            else:
                page_filename = f"page_{i+1:03d}.png"
                page_path = os.path.join(pages_dir, page_filename)
                cv2.imwrite(page_path, img)
                image_files.append(page_path)
                print(f"  Rendered page {i+1}/{total_pages} → {page_filename} ({w}×{h})")
        
        doc.close()
        print(f"  ✅ {total_pages} PDF pages rendered → {len(image_files)} images at {dpi} DPI")
        
        # Process each rendered page
        print(f"\n{'='*60}")
        print(f"  Processing {len(image_files)} pages...")
        print(f"{'='*60}")
        
        all_results = []
        for img_path in image_files:
            result = self.process_single_image(img_path, run_visual_pass=run_visual_pass)
            all_results.append(result)
        
        # Summary
        successful = sum(1 for r in all_results if r.get('success', False))
        total_time = sum(r.get('processing_time_seconds', 0) for r in all_results)
        
        return {
            'success': True,
            'pdf_path': pdf_path,
            'total_pages': total_pages,
            'total_images': len(image_files),
            'successful': successful,
            'failed': len(image_files) - successful,
            'total_time_seconds': round(total_time, 2),
            'results': all_results
        }
    
    def process_directory(self, image_dir: str = None,
                           run_visual_pass: bool = True) -> Dict:
        """Process all images in a directory (if you have images instead of PDF)."""
        if image_dir is None:
            image_dir = os.path.join(self.config.output_dir, "rendered_pages")
        
        extensions = ['*.png', '*.jpg', '*.jpeg', '*.tiff', '*.bmp']
        image_files = []
        for ext in extensions:
            image_files.extend(glob.glob(os.path.join(image_dir, ext)))
        
        image_files.sort()
        
        if not image_files:
            return {'success': False, 'error': f'No images found in {image_dir}'}
        
        print(f"\n{'='*60}")
        print(f"  Processing {len(image_files)} images from: {image_dir}")
        print(f"{'='*60}")
        
        all_results = []
        for img_path in image_files:
            result = self.process_single_image(img_path, run_visual_pass=run_visual_pass)
            all_results.append(result)
        
        # Summary
        successful = sum(1 for r in all_results if r.get('success', False))
        total_time = sum(r.get('processing_time_seconds', 0) for r in all_results)
        
        return {
            'success': True,
            'total_images': len(image_files),
            'successful': successful,
            'failed': len(image_files) - successful,
            'total_time_seconds': round(total_time, 2),
            'results': all_results
        }
    
    def evaluate_with_ground_truth(self, predictions: List[str],
                                     references: List[str]) -> Dict:
        """
        Evaluate pipeline output against ground truth.
        
        Args:
            predictions: List of predicted texts (one per page)
            references: List of ground truth texts (one per page)
        
        Returns:
            Dict with CER, WER, HCPR metrics
        """
        metrics = self.evaluator.evaluate(predictions, references)
        
        print(f"\n{'='*60}")
        print("  EVALUATION RESULTS")
        print(f"{'='*60}")
        print(f"  Average CER:  {metrics['aggregate']['avg_cer']:.4f} ({metrics['aggregate']['avg_cer']*100:.2f}%)")
        print(f"  Average WER:  {metrics['aggregate']['avg_wer']:.4f} ({metrics['aggregate']['avg_wer']*100:.2f}%)")
        print(f"  Average HCPR: {metrics['aggregate']['avg_hcpr']:.4f} ({metrics['aggregate']['avg_hcpr']*100:.2f}%)")
        print(f"  Average BoW:  {metrics['aggregate']['avg_bow_accuracy']:.4f} ({metrics['aggregate']['avg_bow_accuracy']*100:.2f}%)")
        print(f"  Pages: {metrics['aggregate']['num_pages']}")
        print(f"{'='*60}")
        
        return metrics
    
    def _save_outputs(self, result: Dict, image_path: str,
                      line_images: List[np.ndarray], binarized: np.ndarray):
        """Save pipeline outputs to disk."""
        base_name = os.path.splitext(os.path.basename(image_path))[0]
        page_dir = os.path.join(self.config.output_dir, base_name)
        os.makedirs(page_dir, exist_ok=True)
        
        # Save binarized image
        cv2.imwrite(os.path.join(page_dir, "binarized.png"), binarized)
        
        # Save line crops
        lines_dir = os.path.join(page_dir, "lines")
        os.makedirs(lines_dir, exist_ok=True)
        for i, img in enumerate(line_images):
            cv2.imwrite(os.path.join(lines_dir, f"line_{i:03d}.png"), img)
        
        # Save texts
        with open(os.path.join(page_dir, "raw_text.txt"), "w", encoding="utf-8") as f:
            f.write(result['raw_text'])
        
        with open(os.path.join(page_dir, "corrected_text.txt"), "w", encoding="utf-8") as f:
            f.write(result['corrected_text'])
        
        with open(os.path.join(page_dir, "final_text.txt"), "w", encoding="utf-8") as f:
            f.write(result['final_text'])
        
        # Save full result as JSON
        json_result = {k: v for k, v in result.items() 
                       if k != 'ocr_results_per_line'}  # exclude large nested data
        with open(os.path.join(page_dir, "result.json"), "w", encoding="utf-8") as f:
            json.dump(json_result, f, indent=2, ensure_ascii=False)


# ============================================================================
# SECTION 10: GROUND TRUTH LOADER (from .docx file)
# ============================================================================

class GroundTruthLoader:
    """
    Load ground truth text from a .docx file for evaluation.
    
    Supported formats:
      Mode 1 ("per_paragraph"): Each paragraph = one page's ground truth
      Mode 2 ("page_separator"): Pages separated by a line like "--- Page X ---"
      Mode 3 ("full_text"): Entire docx = single page ground truth
    """
    
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Ground truth file not found: {docx_path}")
    
    def load_full_text(self) -> str:
        """Load the entire docx as a single string."""
        doc = DocxDocument(self.docx_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    
    def load_per_paragraph(self) -> List[str]:
        """
        Load ground truth where each non-empty paragraph = one page.
        Best when your docx has one paragraph per page.
        """
        doc = DocxDocument(self.docx_path)
        pages = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        logger.info(f"  Loaded {len(pages)} ground truth pages (per_paragraph mode)")
        return pages
    
    def load_per_page_break(self) -> List[str]:
        """
        Load ground truth split by page breaks in the docx.
        Each section between page breaks = one page's ground truth.
        """
        doc = DocxDocument(self.docx_path)
        
        pages = []
        current_page = []
        
        for paragraph in doc.paragraphs:
            # Check for page break in paragraph's XML
            has_page_break = False
            for run in paragraph.runs:
                if run._element.xml and 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
                    has_page_break = True
                    break
            
            if has_page_break and current_page:
                pages.append("\n".join(current_page))
                current_page = []
            
            if paragraph.text.strip():
                current_page.append(paragraph.text.strip())
        
        # Don't forget the last page
        if current_page:
            pages.append("\n".join(current_page))
        
        logger.info(f"  Loaded {len(pages)} ground truth pages (per_page_break mode)")
        return pages
    
    def load_with_separator(self, separator: str = "---") -> List[str]:
        """
        Load ground truth where pages are separated by a line containing the separator.
        e.g., lines like '--- Page 1 ---' or '---' split pages.
        """
        doc = DocxDocument(self.docx_path)
        all_text = [p.text for p in doc.paragraphs]
        
        pages = []
        current_page = []
        
        for line in all_text:
            if separator in line and len(line.strip().replace(separator, "").replace(" ", "").replace("Page", "").replace("page", "")) < 10:
                # This is a separator line
                if current_page:
                    pages.append("\n".join(current_page))
                    current_page = []
            elif line.strip():
                current_page.append(line.strip())
        
        if current_page:
            pages.append("\n".join(current_page))
        
        logger.info(f"  Loaded {len(pages)} ground truth pages (separator mode)")
        return pages
    
    def load_auto(self) -> List[str]:
        """
        Auto-detect format and load ground truth.
        Tries: page_break → separator → per_paragraph → full_text.
        """
        logger.info(f"Loading ground truth from: {self.docx_path}")
        
        # Try page breaks first
        pages = self.load_per_page_break()
        if len(pages) > 1:
            logger.info(f"  Auto-detected: page_break mode ({len(pages)} pages)")
            return pages
        
        # Try separator mode
        pages = self.load_with_separator("---")
        if len(pages) > 1:
            logger.info(f"  Auto-detected: separator mode ({len(pages)} pages)")
            return pages
        
        # Try per paragraph
        pages = self.load_per_paragraph()
        if len(pages) > 1:
            logger.info(f"  Auto-detected: per_paragraph mode ({len(pages)} pages)")
            return pages
        
        # Fallback: entire docx as one page
        full = self.load_full_text()
        logger.info(f"  Auto-detected: full_text mode (1 page, {len(full)} chars)")
        return [full]


# ============================================================================
# SECTION 11: KAGGLE EXECUTION
# ============================================================================

if __name__ == "__main__":
    
    # =====================================================
    # STEP 1: CONFIGURE PATHS
    # Update these for your Kaggle dataset!
    # =====================================================
    
    config = PipelineConfig(
        # Detectron2 pre-trained weights (upload as Kaggle dataset)
        detectron2_weights="/kaggle/input/your-model-dataset/model_final (8) (1).pth",
        
        # Input PDF file (your document)
        input_pdf="/kaggle/input/your-pdf-dataset/your_document.pdf",
        
        # Output directory
        output_dir="/kaggle/working/ocr_output/",
        
        # Ground truth docx file path
        ground_truth_docx="/kaggle/input/your-dataset/ground_truth.docx",
        
        # PDF rendering DPI (200 is good balance of quality vs speed)
        pdf_dpi=200,
        
        # TrOCR model (downloads from HuggingFace automatically)
        trocr_primary_model="qantev/trocr-large-spanish",
        trocr_fallback_model="microsoft/trocr-base-printed",
        
        # LLM API keys (set as Kaggle secrets)
        openai_api_key=os.getenv("OPENAI_API_KEY", ""),
        gemini_api_key=os.getenv("GEMINI_API_KEY", ""),
    )
    
    # =====================================================
    # STEP 2: INITIALIZE PIPELINE
    # =====================================================
    
    pipeline = HybridOCRPipeline(config)
    
    # =====================================================
    # STEP 3: PROCESS IMAGES
    # =====================================================
    
    # Process the PDF (renders each page → then runs OCR pipeline on each)
    results = pipeline.process_pdf(
        pdf_path=config.input_pdf,
        run_visual_pass=True,  # Set False to skip Gemini visual check
        dpi=config.pdf_dpi
    )
    
    # =====================================================
    # STEP 4: PRINT RESULTS
    # =====================================================
    
    if results['success']:
        print(f"\n{'='*60}")
        print(f"  PIPELINE COMPLETE")
        print(f"{'='*60}")
        print(f"  Total images:   {results['total_images']}")
        print(f"  Successful:     {results['successful']}")
        print(f"  Failed:         {results['failed']}")
        print(f"  Total time:     {results['total_time_seconds']:.1f}s")
        print(f"{'='*60}")
        
        # Print each page's text
        for r in results['results']:
            if r.get('success'):
                print(f"\n--- {os.path.basename(r['image_path'])} ---")
                print(f"Lines detected: {r['num_lines_detected']}")
                print(f"Lines recognized: {r['num_lines_recognized']}")
                print(f"Embellishments filtered: {r['filtered_embellishments']}")
                print(f"Time: {r['processing_time_seconds']}s")
                print(f"\nFinal Text:\n{r['final_text'][:500]}...")
    
    # =====================================================
    # STEP 5: EVALUATE WITH GROUND TRUTH (.docx)
    # =====================================================
    
    print(f"\n{'='*60}")
    print(f"  EVALUATION WITH GROUND TRUTH")
    print(f"{'='*60}")
    
    if os.path.exists(config.ground_truth_docx):
        try:
            # Load ground truth from .docx
            gt_loader = GroundTruthLoader(config.ground_truth_docx)
            ground_truth_pages = gt_loader.load_auto()
            
            print(f"\n  Ground truth loaded: {len(ground_truth_pages)} pages")
            print(f"  Predictions available: {results.get('successful', 0)} pages")
            
            # Get successful predictions
            predictions = [r['final_text'] for r in results['results'] if r.get('success')]
            
            # Match prediction count to ground truth count
            num_eval = min(len(predictions), len(ground_truth_pages))
            
            if num_eval == 0:
                print("\n  ⚠️ No pages to evaluate!")
            else:
                if len(predictions) != len(ground_truth_pages):
                    print(f"\n  ⚠️ Count mismatch: {len(predictions)} predictions vs {len(ground_truth_pages)} ground truth")
                    print(f"  Evaluating first {num_eval} pages...")
                
                eval_predictions = predictions[:num_eval]
                eval_ground_truth = ground_truth_pages[:num_eval]
                
                # ---- Run evaluation at ALL pipeline stages ----
                
                # Stage 4 output (raw TrOCR, no LLM)
                raw_predictions = [r['raw_text'] for r in results['results'] if r.get('success')][:num_eval]
                
                # Stage 5 Pass 1 output (after LLM correction)
                corrected_predictions = [r['corrected_text'] for r in results['results'] if r.get('success')][:num_eval]
                
                # Stage 5 Pass 2 output (after VLM verification — final)
                final_predictions = eval_predictions
                
                print(f"\n{'─'*60}")
                print("  STAGE-BY-STAGE EVALUATION")
                print(f"{'─'*60}")
                
                # Evaluate raw OCR output
                print("\n  📊 After Stage 4 (Raw TrOCR output):")
                raw_metrics = pipeline.evaluate_with_ground_truth(raw_predictions, eval_ground_truth)
                
                # Evaluate after LLM correction
                print("\n  📊 After Stage 5 Pass 1 (LLM text correction):")
                corrected_metrics = pipeline.evaluate_with_ground_truth(corrected_predictions, eval_ground_truth)
                
                # Evaluate final output
                print("\n  📊 After Stage 5 Pass 2 (VLM visual verification — FINAL):")
                final_metrics = pipeline.evaluate_with_ground_truth(final_predictions, eval_ground_truth)
                
                # ---- Summary comparison ----
                print(f"\n{'='*60}")
                print("  EVALUATION SUMMARY — CER IMPROVEMENT ACROSS STAGES")
                print(f"{'='*60}")
                print(f"  {'Stage':<45} {'CER':>8} {'WER':>8}")
                print(f"  {'─'*45} {'─'*8} {'─'*8}")
                print(f"  {'After TrOCR (raw)':<45} {raw_metrics['aggregate']['avg_cer']*100:>7.2f}% {raw_metrics['aggregate']['avg_wer']*100:>7.2f}%")
                print(f"  {'After LLM correction (GPT-4o/Gemini)':<45} {corrected_metrics['aggregate']['avg_cer']*100:>7.2f}% {corrected_metrics['aggregate']['avg_wer']*100:>7.2f}%")
                print(f"  {'After VLM verification (final)':<45} {final_metrics['aggregate']['avg_cer']*100:>7.2f}% {final_metrics['aggregate']['avg_wer']*100:>7.2f}%")
                print(f"  {'─'*45} {'─'*8} {'─'*8}")
                
                # CER improvement
                raw_cer = raw_metrics['aggregate']['avg_cer'] * 100
                final_cer = final_metrics['aggregate']['avg_cer'] * 100
                improvement = raw_cer - final_cer
                print(f"  {'CER improvement (raw → final)':<45} {improvement:>7.2f}pp")
                print(f"{'='*60}")
                
                # ---- Per-page details ----
                print(f"\n{'─'*60}")
                print("  PER-PAGE CER BREAKDOWN")
                print(f"{'─'*60}")
                print(f"  {'Page':>4}  {'Raw CER':>10}  {'LLM CER':>10}  {'Final CER':>10}")
                print(f"  {'─'*4}  {'─'*10}  {'─'*10}  {'─'*10}")
                
                for i in range(num_eval):
                    raw_c = raw_metrics['per_page'][i]['cer'] * 100
                    cor_c = corrected_metrics['per_page'][i]['cer'] * 100
                    fin_c = final_metrics['per_page'][i]['cer'] * 100
                    print(f"  {i+1:>4}  {raw_c:>9.2f}%  {cor_c:>9.2f}%  {fin_c:>9.2f}%")
                
                # ---- Save evaluation results ----
                eval_output = {
                    'ground_truth_file': config.ground_truth_docx,
                    'num_pages_evaluated': num_eval,
                    'raw_ocr_metrics': raw_metrics['aggregate'],
                    'llm_corrected_metrics': corrected_metrics['aggregate'],
                    'final_metrics': final_metrics['aggregate'],
                    'per_page': {
                        'raw': raw_metrics['per_page'],
                        'corrected': corrected_metrics['per_page'],
                        'final': final_metrics['per_page']
                    }
                }
                eval_path = os.path.join(config.output_dir, "evaluation_results.json")
                with open(eval_path, "w", encoding="utf-8") as f:
                    json.dump(eval_output, f, indent=2, ensure_ascii=False)
                print(f"\n  📁 Evaluation results saved to: {eval_path}")
                
        except Exception as e:
            print(f"\n  ❌ Evaluation failed: {e}")
            import traceback
            traceback.print_exc()
    else:
        print(f"\n  ⚠️ Ground truth file not found: {config.ground_truth_docx}")
        print("  Skipping evaluation. To evaluate, upload your ground truth .docx file.")
    
    print("\n✅ Pipeline execution complete! Check /kaggle/working/ocr_output/ for results.")
